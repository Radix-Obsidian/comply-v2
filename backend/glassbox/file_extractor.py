"""Extract text from uploaded documents for compliance scanning."""

import csv
from io import BytesIO, StringIO

from fastapi import HTTPException, UploadFile

MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 MB

ALLOWED_EXTENSIONS = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".txt": "text/plain",
    ".csv": "text/csv",
    ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
}


async def extract_text(file: UploadFile) -> tuple[str, str]:
    """Read uploaded file and extract text content.

    Returns (extracted_text, detected_filetype).
    Raises HTTPException on unsupported type or oversized file.
    """
    filename = file.filename or ""
    ext = _get_extension(filename)

    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type '{ext}'. Allowed: {', '.join(ALLOWED_EXTENSIONS.keys())}",
        )

    content = await file.read()

    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=413,
            detail=f"File too large ({len(content) / 1024 / 1024:.1f} MB). Maximum is {MAX_FILE_SIZE / 1024 / 1024:.0f} MB.",
        )

    if len(content) == 0:
        raise HTTPException(status_code=400, detail="Uploaded file is empty.")

    extractors = {
        ".pdf": _extract_pdf,
        ".docx": _extract_docx,
        ".xlsx": _extract_xlsx,
        ".csv": _extract_csv,
        ".txt": _extract_txt,
    }

    text = extractors[ext](content)

    if not text or not text.strip():
        raise HTTPException(
            status_code=422,
            detail="Could not extract any text from the uploaded file. The file may be scanned/image-based or empty.",
        )

    return text.strip(), ext


def _get_extension(filename: str) -> str:
    """Get lowercase file extension."""
    dot_idx = filename.rfind(".")
    if dot_idx == -1:
        return ""
    return filename[dot_idx:].lower()


def _extract_pdf(content: bytes) -> str:
    """Extract text from PDF using pypdf."""
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(content))
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text)
    return "\n\n".join(pages)


def _extract_docx(content: bytes) -> str:
    """Extract text from DOCX using python-docx."""
    from docx import Document

    doc = Document(BytesIO(content))
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append("\t".join(cells))
    return "\n".join(paragraphs)


def _extract_xlsx(content: bytes) -> str:
    """Extract text from XLSX using openpyxl."""
    from openpyxl import load_workbook

    wb = load_workbook(BytesIO(content), read_only=True, data_only=True)
    lines = []
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        lines.append(f"--- Sheet: {sheet} ---")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                lines.append("\t".join(cells))
    wb.close()
    return "\n".join(lines)


def _extract_csv(content: bytes) -> str:
    """Extract text from CSV."""
    text = content.decode("utf-8", errors="replace")
    reader = csv.reader(StringIO(text))
    lines = []
    for row in reader:
        if any(cell.strip() for cell in row):
            lines.append("\t".join(row))
    return "\n".join(lines)


def _extract_txt(content: bytes) -> str:
    """Extract text from plain text file."""
    return content.decode("utf-8", errors="replace")
